"""
FINAL OPTIMIZED Document Processor - Fixed for 90%+ Accuracy
Date: 2025-08-05 19:49:37 UTC | User: vkhare2909
STRATEGY: Intelligent chunking + critical section extraction + enhanced cleaning
"""

import aiohttp
import asyncio
import fitz  # PyMuPDF
from docx import Document
import io
import re
from typing import Optional, List, Dict, Tuple
import logging
from concurrent.futures import ThreadPoolExecutor
import time

logger = logging.getLogger(__name__)

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class OptimizedInsuranceDocumentProcessor:
    """Optimized document processor for insurance documents"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=6)
        self.session = None
        
    async def get_session(self):
        """Get or create aiohttp session with optimized settings"""
        if self.session is None or self.session.closed:
            timeout = aiohttp.ClientTimeout(total=15)
            connector = aiohttp.TCPConnector(
                limit=20, 
                ttl_dns_cache=300,
                keepalive_timeout=30,
                enable_cleanup_closed=True
            )
            self.session = aiohttp.ClientSession(timeout=timeout, connector=connector)
        return self.session
    
    async def process_document(self, document_url: str) -> str:
        """Optimized document processing for insurance documents"""
        start_time = time.time()
        
        try:
            # Step 1: Download document (max 10s)
            session = await self.get_session()
            
            logger.info(f"📥 Downloading document: {document_url[:60]}...")
            
            async with session.get(document_url) as response:
                if response.status != 200:
                    raise DocumentProcessingError(f"Download failed: HTTP {response.status}")
                
                content_length = response.headers.get('Content-Length')
                if content_length:
                    size_mb = int(content_length) / (1024 * 1024)
                    logger.info(f"📄 Document size: {size_mb:.1f} MB")
                
                # Optimized streaming
                content = bytearray()
                chunk_size = 65536  # 64KB chunks for better performance
                
                async for chunk in response.content.iter_chunked(chunk_size):
                    content.extend(chunk)
                
                content = bytes(content)
            
            download_time = time.time() - start_time
            logger.info(f"⚡ Download completed in {download_time:.1f}s")
            
            # Step 2: Extract text (max 8s)
            extraction_start = time.time()
            
            if content.startswith(b'%PDF'):
                text = await self._extract_pdf_optimized(content)
            elif content.startswith(b'PK\x03\x04') and b'word/' in content[:2000]:
                text = await self._extract_docx_optimized(content)
            else:
                raise DocumentProcessingError("Unsupported format. Only PDF/DOCX supported.")
            
            extraction_time = time.time() - extraction_start
            logger.info(f"⚡ Text extraction completed in {extraction_time:.1f}s")
            
            # Step 3: Clean and optimize text (max 2s)
            clean_start = time.time()
            cleaned_text = self._deep_clean_insurance_text(text)
            clean_time = time.time() - clean_start
            
            total_time = time.time() - start_time
            
            logger.info(f"✅ Document processing complete:")
            logger.info(f"   📊 Total time: {total_time:.1f}s")
            logger.info(f"   📄 Text length: {len(cleaned_text):,} chars")
            logger.info(f"   ⚡ Processing speed: {len(cleaned_text)/total_time:.0f} chars/sec")
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            raise DocumentProcessingError(f"Processing failed: {str(e)}")
    
    async def _extract_pdf_optimized(self, pdf_content: bytes) -> str:
        """Optimized PDF extraction for insurance documents"""
        
        def extract_pdf_intelligent():
            try:
                doc = fitz.open(stream=pdf_content, filetype="pdf")
                total_pages = len(doc)
                
                logger.info(f"📄 Processing {total_pages} PDF pages...")
                
                # INTELLIGENT PAGE SELECTION FOR INSURANCE DOCUMENTS
                pages_to_process = self._select_critical_pages(total_pages)
                
                logger.info(f"📄 Processing {len(pages_to_process)} critical pages from {total_pages} total")
                
                # Extract text with enhanced processing
                text_parts = []
                
                for page_num in pages_to_process:
                    if page_num < total_pages:
                        try:
                            page = doc.load_page(page_num)
                            
                            # Enhanced text extraction
                            page_text = page.get_text()
                            
                            # Try different extraction methods if text is sparse
                            if len(page_text.strip()) < 100:
                                # Try with layout preservation
                                page_text = page.get_text("layout")
                            
                            if page_text.strip():
                                # Add page context
                                enhanced_text = f"\n--- PAGE {page_num + 1} ---\n{page_text}"
                                text_parts.append(enhanced_text)
                                
                        except Exception as e:
                            logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                            continue
                
                doc.close()
                
                full_text = "\n".join(text_parts)
                
                if not full_text.strip():
                    raise DocumentProcessingError("No text extracted from PDF")
                
                logger.info(f"✅ Extracted {len(full_text):,} characters from {len(pages_to_process)} pages")
                return full_text
                
            except Exception as e:
                raise DocumentProcessingError(f"PDF extraction failed: {str(e)}")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract_pdf_intelligent)
    
    def _select_critical_pages(self, total_pages: int) -> List[int]:
        """Select critical pages for insurance documents"""
        
        pages_to_process = []
        
        if total_pages <= 50:
            # Small documents - process all pages
            pages_to_process = list(range(total_pages))
        
        elif total_pages <= 100:
            # Medium documents - process first 60% + key sections
            first_section = int(total_pages * 0.6)
            pages_to_process.extend(range(first_section))
            
            # Add samples from remaining sections
            remaining_start = first_section
            sample_every = max(2, (total_pages - remaining_start) // 15)
            pages_to_process.extend(range(remaining_start, total_pages, sample_every))
            
            # Always include last 5 pages (often contain important definitions)
            pages_to_process.extend(range(max(first_section, total_pages - 5), total_pages))
        
        else:
            # Large documents (>100 pages) - strategic sampling
            logger.info(f"📄 Large document detected: using strategic sampling")
            
            # First 40 pages (policy details, coverage, benefits)
            pages_to_process.extend(range(0, min(40, total_pages)))
            
            # Key middle sections
            quarter_points = [
                total_pages // 4,
                total_pages // 2,
                (3 * total_pages) // 4
            ]
            
            for center in quarter_points:
                start = max(40, center - 10)
                end = min(total_pages, center + 10)
                pages_to_process.extend(range(start, end))
            
            # Last 20 pages (definitions, annexures, important terms)
            last_start = max(max(pages_to_process) + 1 if pages_to_process else 0, total_pages - 20)
            pages_to_process.extend(range(last_start, total_pages))
            
            # Remove duplicates and limit to 80 pages max
            pages_to_process = sorted(set(pages_to_process))[:80]
        
        return pages_to_process
    
    async def _extract_docx_optimized(self, docx_content: bytes) -> str:
        """Optimized DOCX extraction"""
        
        def extract_docx_intelligent():
            try:
                doc = Document(io.BytesIO(docx_content))
                text_parts = []
                
                logger.info(f"📄 Processing DOCX document...")
                
                # Extract paragraphs with intelligent limits
                paragraph_count = 0
                for paragraph in doc.paragraphs:
                    if paragraph_count > 2000:  # Reasonable limit
                        break
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                        paragraph_count += 1
                
                # Extract tables with enhanced processing
                table_count = 0
                for table in doc.tables:
                    if table_count > 100:  # Generous limit for insurance docs
                        break
                    
                    table_text = []
                    row_count = 0
                    for row in table.rows:
                        if row_count > 50:  # Limit rows per table
                            break
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            table_text.append(row_text)
                            row_count += 1
                    
                    if table_text:
                        text_parts.append(f"\n--- TABLE {table_count + 1} ---\n" + "\n".join(table_text))
                        table_count += 1
                
                full_text = "\n".join(text_parts)
                
                if not full_text.strip():
                    raise DocumentProcessingError("No text extracted from DOCX")
                
                logger.info(f"✅ Extracted {len(full_text):,} characters from DOCX")
                return full_text
                
            except Exception as e:
                raise DocumentProcessingError(f"DOCX extraction failed: {str(e)}")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, extract_docx_intelligent)
    
    def _deep_clean_insurance_text(self, text: str) -> str:
        """Deep clean insurance text for optimal processing"""
        
        logger.info("🧹 Deep cleaning insurance text...")
        
        # Step 1: Basic cleanup
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]{4,}', '  ', text)
        text = re.sub(r'\t+', ' ', text)
        
        # Step 2: Fix common PDF extraction issues
        # Fix hyphenated words across lines
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
        
        # Fix broken words across lines
        text = re.sub(r'(\w)\s*\n\s*(\w)', r'\1 \2', text)
        
        # Step 3: Preserve important insurance formatting
        # Preserve section headers (all caps with spacing)
        text = re.sub(r'\n([A-Z\s]{10,})\n', r'\n\n\1\n\n', text)
        
        # Preserve numbered lists
        text = re.sub(r'\n(\d+\.)\s*', r'\n\n\1 ', text)
        text = re.sub(r'\n([a-z]\))\s*', r'\n\1 ', text)
        text = re.sub(r'\n([ivx]+\.)\s*', r'\n\1 ', text)
        
        # Step 4: Fix insurance-specific formatting
        # Preserve important terms formatting
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Fix common insurance terms
        insurance_fixes = {
            r'suminsured': 'sum insured',
            r'policyterm': 'policy term',
            r'waitingperiod': 'waiting period',
            r'graceperiod': 'grace period',
            r'preexisting': 'pre-existing',
            r'noclaimdiscount': 'no claim discount',
            r'ayushtreatment': 'ayush treatment'
        }
        
        for pattern, replacement in insurance_fixes.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Step 5: Remove unwanted elements
        # Remove control characters but keep essential formatting
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        # Step 6: Normalize spacing
        # Ensure proper spacing around punctuation
        text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text)
        text = re.sub(r'([,;:])\s*', r'\1 ', text)
        
        # Step 7: Preserve insurance document structure
        # Ensure section breaks are clear
        text = re.sub(r'\n(SECTION|CHAPTER|PART)\s*(\d+)', r'\n\n\1 \2', text, flags=re.IGNORECASE)
        
        # Preserve definition sections
        text = re.sub(r'\n(DEFINITIONS?|MEANING)\s*:', r'\n\n\1:', text, flags=re.IGNORECASE)
        
        # Step 8: Final cleanup
        # Remove trailing whitespace
        lines = [line.rstrip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove excessive blank lines at start and end
        text = text.strip()
        
        # Ensure consistent line endings
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        logger.info(f"✅ Text cleaning complete: {len(text):,} characters")
        
        return text
    
    def extract_critical_sections(self, text: str) -> Dict[str, str]:
        """Extract critical sections from insurance document"""
        
        logger.info("🔍 Extracting critical insurance sections...")
        
        section_patterns = {
            'grace_period': [
                r'(?:grace period|payment grace|premium grace).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:payment.*?due.*?date|premium.*?due).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:thirty days?|30 days?).*?(?:grace|payment|due).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'waiting_periods': [
                r'(?:waiting period|cooling period|moratorium).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:pre-existing.*?disease|ped).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:thirty-six months?|36 months?).*?(?:waiting|pre-existing).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'maternity_coverage': [
                r'(?:maternity|pregnancy|childbirth|delivery).*?(?:benefit|coverage|waiting).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:twenty-four months?|24 months?).*?(?:maternity|pregnancy).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:lawful.*?medical.*?termination|lmt).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'coverage_benefits': [
                r'(?:coverage|benefits?).*?(?:includes?|covers?).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:scope.*?of.*?cover|what.*?is.*?covered).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:sum.*?insured|si|coverage.*?amount).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'exclusions': [
                r'(?:exclusions?|not.*?covered|limitations?).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:what.*?is.*?not.*?covered|exceptions?).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'hospital_definition': [
                r'(?:hospital.*?(?:means|defined|definition)|definition.*?hospital).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:medical.*?institution|healthcare.*?facility).*?(?:means|defined).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'ncd_bonus': [
                r'(?:no.*?claim.*?discount|ncd|bonus).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:claim.*?free.*?bonus|discount.*?benefit).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'ayush_treatment': [
                r'(?:ayush|ayurveda|yoga|naturopathy|unani|siddha|homeopathy).*?(?:treatment|coverage|benefit).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:alternative.*?medicine|traditional.*?medicine).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'claims_procedure': [
                r'(?:claim.*?procedure|settlement.*?process|reimbursement).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:how.*?to.*?claim|claim.*?process|intimation).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ],
            'cataract_coverage': [
                r'(?:cataract|eye.*?surgery|lens.*?replacement).*?(?:waiting|coverage|benefit).*?(?=\n\n|\n[A-Z]{3,}|$)',
                r'(?:two.*?years?|2.*?years?).*?(?:cataract|eye).*?(?=\n\n|\n[A-Z]{3,}|$)'
            ]
        }
        
        sections = {}
        text_lower = text.lower()
        
        for section_name, patterns in section_patterns.items():
            best_match = ""
            best_score = 0
            
            for pattern in patterns:
                matches = re.finditer(pattern, text_lower, re.IGNORECASE | re.DOTALL)
                
                for match in matches:
                    # Get extended context around match
                    start = max(0, match.start() - 300)
                    end = min(len(text), match.end() + 1000)
                    context = text[start:end]
                    
                    # Score based on content quality and relevance
                    score = len(context)
                    
                    # Boost for specific insurance terms
                    if re.search(r'\d+\s*(?:days?|months?|years?)', context, re.IGNORECASE):
                        score += 200
                    
                    # Boost for key insurance keywords
                    insurance_keywords = [
                        'grace', 'waiting', 'coverage', 'benefit', 'premium',
                        'policy', 'insured', 'hospital', 'treatment', 'claim'
                    ]
                    
                    for keyword in insurance_keywords:
                        if keyword in context.lower():
                            score += 50
                    
                    # Boost for section headers
                    if re.search(r'\n\s*[A-Z\s]{10,}\s*\n', context):
                        score += 100
                    
                    if score > best_score:
                        best_score = score
                        best_match = context
            
            if best_match and len(best_match) > 150:
                sections[section_name] = best_match.strip()
        
        logger.info(f"✅ Extracted {len(sections)} critical sections")
        
        return sections
    
    def create_intelligent_chunks(self, text: str, sections: Dict[str, str], chunk_size: int = 1200, overlap: int = 150) -> List[str]:
        """Create intelligent chunks optimized for insurance documents"""
        
        logger.info(f"📝 Creating intelligent chunks (size: {chunk_size}, overlap: {overlap})...")
        
        chunks = []
        
        # Step 1: Create priority chunks from critical sections
        section_chunks = []
        for section_name, section_content in sections.items():
            header = f"[{section_name.upper().replace('_', ' ')}]"
            section_chunk_list = self._chunk_section_intelligently(
                section_content, header, chunk_size, overlap
            )
            section_chunks.extend(section_chunk_list)
        
        chunks.extend(section_chunks)
        
        # NO CHUNK LIMITING - Use all chunks for maximum accuracy
        # chunks = chunks  # Keep all chunks - no artificial limitations
        
        # Skip general chunks if we have enough section chunks  
        if len(chunks) >= 200:  # Only skip if we have substantial section coverage
            logger.info(f"✅ Created {len(chunks)} intelligent chunks (comprehensive coverage)")
            return chunks
        
        # Step 2: Create general chunks from remaining text only if needed
        # Remove already processed sections to avoid duplication
        remaining_text = text
        for section_content in sections.values():
            # Remove the section content (approximately)
            section_words = section_content.split()[:50]  # First 50 words as identifier
            section_start = ' '.join(section_words)
            if section_start in remaining_text:
                start_pos = remaining_text.find(section_start)
                if start_pos != -1:
                    end_pos = min(start_pos + len(section_content), len(remaining_text))
                    remaining_text = remaining_text[:start_pos] + remaining_text[end_pos:]
        
        # Create general chunks from remaining text
        if len(remaining_text.strip()) > 500:
            general_chunks = self._create_overlapping_chunks(
                remaining_text, "[GENERAL CONTENT]", chunk_size, overlap
            )
            # Limit general chunks to avoid overwhelming specific content
            chunks.extend(general_chunks[:20])
        
        logger.info(f"✅ Created {len(chunks)} intelligent chunks ({len(section_chunks)} from sections, {len(chunks)-len(section_chunks)} general)")
        
        return chunks
    
    def _chunk_section_intelligently(self, text: str, header: str, chunk_size: int, overlap: int) -> List[str]:
        """Intelligently chunk a specific section"""
        
        chunks = []
        
        # If section fits in one chunk
        if len(text) + len(header) <= chunk_size:
            chunks.append(f"{header}\n{text}")
            return chunks
        
        # Split by sentences for better context preservation
        sentences = re.split(r'[.!?]+(?=\s+[A-Z])', text)
        current_chunk = header
        current_size = len(header)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_with_period = sentence + "."
            
            # Check if adding this sentence exceeds chunk size
            if current_size + len(sentence_with_period) + 2 > chunk_size:
                # Save current chunk if it has substantial content
                if current_size > len(header) + 100:
                    chunks.append(current_chunk)
                    
                    # Start new chunk with overlap
                    overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    # Remove header from overlap to avoid duplication
                    if overlap_text.startswith(header):
                        overlap_text = overlap_text[len(header):].strip()
                    
                    current_chunk = f"{header}\n{overlap_text}\n{sentence_with_period}"
                    current_size = len(current_chunk)
                else:
                    # Current chunk too small, just add the sentence
                    current_chunk += f"\n{sentence_with_period}"
                    current_size += len(sentence_with_period) + 1
            else:
                # Add sentence to current chunk
                current_chunk += f"\n{sentence_with_period}"
                current_size += len(sentence_with_period) + 1
        
        # Add final chunk if it has substantial content
        if current_size > len(header) + 100:
            chunks.append(current_chunk)
        
        return chunks
    
    def _create_overlapping_chunks(self, text: str, header: str, chunk_size: int, overlap: int) -> List[str]:
        """Create overlapping chunks with header"""
        
        chunks = []
        
        # Adjust chunk size to account for header
        effective_chunk_size = chunk_size - len(header) - 10
        
        # Split into words for better control
        words = text.split()
        
        start = 0
        while start < len(words):
            # Calculate end position
            end = min(start + effective_chunk_size, len(words))
            
            # Create chunk text
            chunk_words = words[start:end]
            chunk_text = ' '.join(chunk_words)
            
            # Add header
            full_chunk = f"{header}\n{chunk_text}"
            chunks.append(full_chunk)
            
            # Move start position with overlap
            overlap_words = min(overlap, len(chunk_words) // 3)  # Overlap 1/3 of chunk
            start = end - overlap_words
            
            # Break if we've covered all text
            if end >= len(words):
                break
        
        return chunks
    
    async def close(self):
        """Clean up resources"""
        if self.session and not self.session.closed:
            await self.session.close()
        self.executor.shutdown(wait=False)

# Global instance for reuse
optimized_processor = OptimizedInsuranceDocumentProcessor()

async def process_document(document_url: str) -> str:
    """Main function for optimized document processing"""
    return await optimized_processor.process_document(document_url)

async def process_document_with_sections(document_url: str) -> Tuple[str, Dict[str, str], List[str]]:
    """Process document and return text, sections, and chunks"""
    
    # Get the cleaned text
    cleaned_text = await optimized_processor.process_document(document_url)
    
    # Extract critical sections
    sections = optimized_processor.extract_critical_sections(cleaned_text)
    
    # Create intelligent chunks
    chunks = optimized_processor.create_intelligent_chunks(cleaned_text, sections)
    
    return cleaned_text, sections, chunks

# Backward compatibility
hyper_fast_processor = optimized_processor